import os
from datetime import datetime
import asyncio
from langchain_openai import ChatOpenAI
from app.agents.schema.resume_create_agent import ResumeAgentState
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from markdown2 import markdown
from app.agents.base_node import LLMBaseNode


class CreateResumeNode(LLMBaseNode):
    def __init__(self, llm: ChatOpenAI):
        super().__init__(llm)
        # 환경 설정
        self.environment = os.getenv(
            "ENVIRONMENT", "development"
        )  # development, production
        self.use_s3 = os.getenv("USE_S3", "false").lower() == "true"

        self.logger.info(f"환경: {self.environment}, S3 사용: {self.use_s3}")

    async def execute(self, state: ResumeAgentState) -> ResumeAgentState:
        try:
            prompt = self._build_resume_prompt(state)

            # LLM으로 이력서 생성
            content = await self._generate_resume_content(prompt)

            # 생성된 내용을 임시 저장 (에러 처리용)
            self._last_generated_content = content

            # 환경에 따른 파일 처리
            if self.use_s3:
                # 프로덕션: S3 업로드
                docx_path = await self._create_and_upload_document(state, content)
            else:
                # 개발: 로컬 저장
                docx_path = await self._create_local_document(state, content)

            # 상태 업데이트
            state.docx_path = docx_path
            state.resume = content
            state.step = "completed"

            return state

        except Exception as e:
            self.logger.error(f"CreateResumeNode 실행 중 오류: {e}")

            # 에러가 발생해도 기본적인 이력서 내용은 제공
            if (
                hasattr(self, "_last_generated_content")
                and self._last_generated_content
            ):
                state.resume = self._last_generated_content
            else:
                # LLM 내용도 없다면 기본 이력서 생성
                state.resume = self._create_fallback_resume(state)

            # 에러 상태 설정
            state.step = "completed_with_error"
            state.docx_path = ""  # 파일 생성 실패

            return state

    async def _create_and_upload_document(
        self, state: ResumeAgentState, content: str
    ) -> str:
        """Word 문서 생성 후 S3 업로드 (프로덕션용)"""

        # 1. 임시 로컬 파일 생성
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        local_filename = f"resume_agent_{timestamp}.docx"
        temp_dir = "/tmp" if os.path.exists("/tmp") else "temp"
        os.makedirs(temp_dir, exist_ok=True)
        local_path = os.path.join(temp_dir, local_filename)

        # 2. Word 문서 생성
        def create_doc():
            doc = Document()
            doc.add_heading("이력서", level=0).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # 기본 정보 섹션
            self._add_basic_info_section(doc, state.inputs)

            # LLM 생성 내용 섹션
            self._render_llm_content_stylized(content, doc)

            doc.save(local_path)

        await asyncio.to_thread(create_doc)

        # 3. S3 업로드
        try:
            from app.utils.upload_file_to_s3 import async_upload_file_to_s3
            from io import BytesIO

            # 파일을 BytesIO로 읽기
            def read_file_to_bytes():
                with open(local_path, "rb") as f:
                    return f.read()

            file_bytes = await asyncio.to_thread(read_file_to_bytes)
            file_obj = BytesIO(file_bytes)
            file_obj.name = local_filename

            s3_url = await async_upload_file_to_s3(file_obj, local_filename)

            if not s3_url:
                self.logger.error("S3 업로드 실패")
                raise Exception("S3 업로드에 실패했습니다")

            self.logger.info(f"S3 업로드 성공: {s3_url}")

            # 4. 임시 파일 정리
            try:
                os.remove(local_path)
                self.logger.info(f"임시 파일 정리 완료: {local_path}")
            except Exception as e:
                self.logger.warning(f"임시 파일 정리 실패: {e}")

            return s3_url

        except Exception as e:
            self.logger.error(f"S3 업로드 중 오류: {e}")
            # S3 업로드 실패시 로컬 경로 반환 (fallback)
            return local_path

    async def _create_local_document(
        self, state: ResumeAgentState, content: str
    ) -> str:
        """Word 문서 로컬 생성 (개발용)"""

        # 절대 경로 생성
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"resume_final_{timestamp}.docx"

        # 프로젝트 루트의 generated 폴더에 저장
        current_dir = os.getcwd()
        abs_dir = os.path.join(current_dir, "generated")
        abs_path = os.path.join(abs_dir, filename)

        def write_to_doc():
            try:
                # 디렉토리 생성
                os.makedirs(abs_dir, exist_ok=True)

                # 문서 생성
                doc = Document()
                doc.add_heading("이력서", level=0).alignment = (
                    WD_PARAGRAPH_ALIGNMENT.LEFT
                )

                # 기본 정보 섹션
                self._add_basic_info_section(doc, state.inputs)

                # LLM 생성 내용 섹션
                self._render_llm_content_stylized(content, doc)

                # 파일 저장 (abs_path 사용)
                doc.save(abs_path)

                # 파일 생성 확인
                if os.path.exists(abs_path):
                    file_size = os.path.getsize(abs_path)
                    print(
                        f"✅ 로컬 파일 생성 성공: {abs_path} (크기: {file_size} bytes)"
                    )
                else:
                    print(f"❌ 로컬 파일 생성 실패: {abs_path}")

            except Exception as e:
                print(f"❌ 문서 생성 중 오류: {e}")
                raise

        await asyncio.to_thread(write_to_doc)

        # 생성 후 다시 한번 확인
        if os.path.exists(abs_path):
            self.logger.info(f"로컬 이력서 파일 생성 완료: {abs_path}")
            return abs_path
        else:
            self.logger.error(f"로컬 파일 생성 실패: {abs_path}")
            raise FileNotFoundError(f"파일을 생성할 수 없습니다: {abs_path}")

    def _build_resume_prompt(self, state: ResumeAgentState) -> str:
        # LLM에 보낼 요약 정보
        base_info = f"""
    이메일: {state.inputs.email}
    희망 직무: {state.inputs.preferred_job}
    전공 여부: {state.inputs.major_type}
    재직 회사: {state.inputs.company_name}
    직무명: {state.inputs.position}
    재직 기간: {state.inputs.work_period}개월
    자격증 수: {state.inputs.certification_count}
    프로젝트 수: {state.inputs.project_count}
    추가 경험: {state.inputs.additional_experiences}
        """

        qna_info = "\n".join(
            [f"Q: {a['question']}\nA: {a['answer']}" for a in state.answers]
        )

        return f"""
    다음은 이력서에 포함될 정보입니다. 아래 정보를 기반으로 고급 이력서 초안을 마크다운 형식으로 작성해주세요. 항목: 경력 사항, 프로젝트, 기술 역량, 자격증 등

    [입력 정보]
    {base_info}

    [질문 응답]
    {qna_info}
    """

    async def _generate_resume_content(self, prompt: str) -> str:
        """LLM으로 이력서 내용 생성"""
        content = await self._safe_llm_call(
            prompt, "이력서 생성 중 오류가 발생했습니다."
        )
        return content

    def _add_basic_info_section(self, doc: Document, inputs):
        """기본 정보 섹션 추가"""
        doc.add_heading("기본 정보", level=1)
        doc.add_paragraph(f"이메일: {inputs.email}")
        doc.add_paragraph(f"희망 직무: {inputs.preferred_job}")
        doc.add_paragraph(f"전공 여부: {inputs.major_type}")

    def _render_llm_content_stylized(self, markdown_text: str, doc: Document):
        """마크다운 텍스트를 Word 문서에 스타일 적용하여 렌더링"""
        html = markdown(markdown_text)
        soup = BeautifulSoup(html, "html.parser")

        for element in soup.descendants:
            if element.name == "h1":
                section = doc.add_heading(element.get_text(), level=1)
                self._add_horizontal_line(section)
            elif element.name == "h2":
                doc.add_heading(element.get_text(), level=2)
            elif element.name == "p":
                doc.add_paragraph(element.get_text())
            elif element.name == "li":
                doc.add_paragraph("▪ " + element.get_text())

    def _create_fallback_resume(self, state: ResumeAgentState) -> str:
        """파일 생성 실패시 기본 이력서 텍스트 생성"""

        inputs = state.inputs
        answers = state.answers

        fallback_content = f"""# 이력서

## 기본 정보
- **이메일**: {inputs.email}
- **희망 직무**: {inputs.preferred_job}
- **전공 여부**: {inputs.major_type}
- **재직 회사**: {inputs.company_name}
- **현재 직무**: {inputs.position}
- **재직 기간**: {inputs.work_period}개월
- **자격증 수**: {inputs.certification_count}개
- **프로젝트 수**: {inputs.project_count}개

## 추가 경험
{inputs.additional_experiences}

## 상세 정보
"""

        for i, qa in enumerate(answers, 1):
            fallback_content += f"\n### Q{i}: {qa['question']}\n{qa['answer']}\n"

        fallback_content += f"\n\n---\n⚠️ 파일 생성 중 오류가 발생하여 텍스트 형태로 제공됩니다.\n텍스트를 복사하여 별도 문서로 저장해주세요."

        return fallback_content

    def _add_horizontal_line(self, paragraph):
        """단락에 수평선 추가"""
        pPr = paragraph._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "2F5496")
        border.append(bottom)
        pPr.append(border)
