import os
import logging
import traceback
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO
from datetime import datetime
import asyncio
from app.schemas.resume_create import ResumeCreateRequest
from app.agents.schema.resume_create_agent import ResumeAgentState
from app.agents.nodes.create_resume import CreateResumeNode
from app.utils.upload_file_to_s3 import upload_file_to_s3
from langchain_openai import ChatOpenAI

# 로깅 설정
logging.basicConfig(
    level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
)
logger = logging.getLogger(__name__)

# LLM 초기화
llm = ChatOpenAI(model="gpt-3.5-turbo", temperature=0.3)


# ------------------------- 스타일 관련 도우미 함수들 -------------------------
# 구분선 넣기 함수
def add_horizontal_line(paragraph, before=6, after=6):
    """구분선 넣기 함수"""
    try:
        paragraph.paragraph_format.space_before = Pt(before)
        paragraph.paragraph_format.space_after = Pt(after)

        pPr = paragraph._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")  # 실선
        bottom.set(qn("w:sz"), "8")  # 굵기 (2~24까지 조절)
        bottom.set(qn("w:space"), "0")  # 줄과 간격
        bottom.set(qn("w:color"), "2F5496")  # 색상
        border.append(bottom)
        pPr.append(border)
    except Exception as e:
        logger.error(f"구분선 추가 실패: {e}")


# 밑줄 있는 빈칸 라벨 생성 헬퍼
def add_underlined_paragraph(doc, label: str, length: int = 40):
    """밑줄 있는 빈칸 라벨 생성 헬퍼"""
    try:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ")
        run = p.add_run(" " * length)
        run.font.underline = True
        run.font.size = Pt(11)
    except Exception as e:
        logger.error(f"밑줄 단락 추가 실패: {e}")


# ------------------------- 이력서 생성 핵심 로직 -------------------------
def _generate_resume_doc(data: ResumeCreateRequest) -> BytesIO:
    """기본 이력서 템플릿 생성"""
    try:
        logger.info("이력서 문서 생성 시작")
        doc = Document()

        # 제목
        title = doc.add_heading("OOO 이력서", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 기본 정보
        doc.add_heading("기본 정보", level=1)
        doc.add_paragraph(f"지원 분야: {data.preferred_job}")
        doc.add_paragraph("생년월일:")
        doc.add_paragraph(f"이메일: {data.email}")
        doc.add_paragraph("전화번호: ")
        doc.add_paragraph("GitHub URL:")

        if data.major_type == "MAJOR":
            doc.add_paragraph("전공 : 컴퓨터 공학 관련 전공")
        else:
            doc.add_paragraph("전공 : ")

        doc.add_paragraph(" " + " " * 100)

        # 경력 사항
        if data.work_period is not None and data.work_period > 0:
            heading_paragraph = doc.add_heading("경력 사항", level=1)
            add_horizontal_line(heading_paragraph)

            if data.work_period < 12:
                period_text = f"{data.work_period}개월"
            else:
                years = data.work_period // 12
                months = data.work_period % 12
                period_text = f"{years}년 {months}개월" if months > 0 else f"{years}년"

            add_underlined_paragraph(doc, f"회사명: {data.company_name}")
            if data.position:
                doc.add_paragraph(f"직무명: {data.position}")
            add_underlined_paragraph(doc, f"근무기간: {period_text}")
            doc.add_paragraph("담당 업무:")
            doc.add_paragraph("▪ " + " " * 100)
            doc.add_paragraph("▪ " + " " * 100)

        # 프로젝트 경험
        if data.project_count > 0:
            heading_paragraph = doc.add_heading("프로젝트", level=1)
            add_horizontal_line(heading_paragraph)

            for i in range(data.project_count):
                doc.add_paragraph(f"[프로젝트 {i + 1} 제목]")
                add_underlined_paragraph(doc, "기간", 40)
                doc.add_paragraph("프로젝트 요약:")
                doc.add_paragraph("▪ " + " " * 100)
                doc.add_paragraph("맡은 역할:")
                doc.add_paragraph("▪ " + " " * 100)
                doc.add_paragraph("▪ " + " " * 100)
                doc.add_paragraph("성과:")
                doc.add_paragraph("▪ " + " " * 100)

        # 자격증
        if data.certification_count > 0:
            heading_paragraph = doc.add_heading("자격증", level=1)
            add_horizontal_line(heading_paragraph)

            for i in range(data.certification_count):
                add_underlined_paragraph(doc, f"▪ ", 50)

        # 기타
        if data.additional_experiences:
            heading_paragraph = doc.add_heading("기타", level=1)
            add_horizontal_line(heading_paragraph)
            doc.add_paragraph(f"{data.additional_experiences}")

        # BytesIO로 저장
        byte_io = BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)

        logger.info("이력서 문서 생성 완료")
        return byte_io

    except Exception as e:
        logger.error(f"이력서 문서 생성 실패: {e}")
        logger.error(traceback.format_exc())
        raise


# ------------------------- 외부 호출용 async 함수들 -------------------------
async def generate_resume_draft(data: ResumeCreateRequest) -> str:
    """기본 이력서 초안 생성 및 S3 업로드"""
    try:
        logger.info("이력서 초안 생성 시작")

        # 파일명 생성
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"resume_draft_{timestamp}.docx"

        # docx 생성 (비동기 스레드 오프로드)
        byte_stream = await asyncio.to_thread(_generate_resume_doc, data)

        # S3 업로드
        s3_file_path = f"resume/{filename}"
        file_url = await asyncio.to_thread(upload_file_to_s3, byte_stream, s3_file_path)

        logger.info(f"이력서 초안 생성 및 업로드 완료: {file_url}")
        return file_url

    except Exception as e:
        logger.error(f"이력서 초안 생성 실패: {e}")
        logger.error(traceback.format_exc())
        raise


async def save_agent_resume_to_docx(state: ResumeAgentState) -> str:
    """에이전트가 생성한 이력서를 DOCX 파일로 저장"""
    try:
        logger.info("에이전트 이력서 저장 시작")

        # 파일명 및 경로 생성
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"resume_agent_{timestamp}.docx"
        save_dir = "./generated_resumes"
        os.makedirs(save_dir, exist_ok=True)
        filepath = os.path.join(save_dir, filename)

        # LLM 결과 생성 + 저장 로직
        def _write_to_docx():
            try:
                # CreateResumeNode 인스턴스 생성 및 실행
                resume_node = CreateResumeNode(llm)
                updated_state = resume_node.execute(state)

                # 이력서 텍스트 추출
                resume_text = ""
                if hasattr(updated_state, "resume"):
                    resume_text = updated_state.resume
                elif isinstance(updated_state, dict):
                    resume_text = updated_state.get("resume", "")
                else:
                    logger.warning("이력서 텍스트를 찾을 수 없습니다.")
                    resume_text = "이력서 생성에 실패했습니다."

                # DOCX 문서 생성
                doc = Document()

                if resume_text:
                    # 줄바꿈으로 분리하여 각 줄을 단락으로 추가
                    lines = resume_text.split("\n")
                    for line in lines:
                        line = line.strip()
                        if line:  # 빈 줄이 아닌 경우만 추가
                            # 제목인지 확인 (간단한 휴리스틱)
                            if line.startswith("#") or line.isupper() or len(line) < 50:
                                # 제목으로 처리
                                heading = doc.add_heading(
                                    line.replace("#", "").strip(), level=1
                                )
                            else:
                                # 일반 단락으로 처리
                                doc.add_paragraph(line)
                        else:
                            # 빈 줄 추가
                            doc.add_paragraph("")
                else:
                    doc.add_paragraph("이력서 내용이 생성되지 않았습니다.")

                # 파일 저장
                doc.save(filepath)
                logger.info(f"이력서 파일 저장 완료: {filepath}")

            except Exception as e:
                logger.error(f"DOCX 생성 중 오류: {e}")
                logger.error(traceback.format_exc())

                # 오류 발생 시 기본 문서 생성
                doc = Document()
                doc.add_heading("이력서 생성 오류", level=1)
                doc.add_paragraph(f"이력서 생성 중 오류가 발생했습니다: {str(e)}")
                doc.save(filepath)
                raise

        # 비동기 스레드에서 실행
        await asyncio.to_thread(_write_to_docx)

        logger.info(f"에이전트 이력서 저장 완료: {filepath}")
        return filepath

    except Exception as e:
        logger.error(f"에이전트 이력서 저장 실패: {e}")
        logger.error(traceback.format_exc())
        raise


def validate_resume_data(data: ResumeCreateRequest) -> bool:
    """이력서 데이터 유효성 검증"""
    try:
        if not data.email or "@" not in data.email:
            logger.error("유효하지 않은 이메일 주소")
            return False

        if not data.preferred_job:
            logger.error("선호 직무가 비어있습니다")
            return False

        if data.project_count < 0 or data.certification_count < 0:
            logger.error("프로젝트 또는 자격증 개수가 음수입니다")
            return False

        return True

    except Exception as e:
        logger.error(f"데이터 검증 중 오류: {e}")
        return False
